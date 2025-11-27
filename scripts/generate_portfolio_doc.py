"""
Generate a Word (.docx) document from the Markdown summary in `docs/Portfolio_Functionality.md`.
This script uses `python-docx` (install via `pip install python-docx`).
"""
import sys
from pathlib import Path

MD = Path(__file__).resolve().parents[1] / 'docs' / 'Portfolio_Functionality.md'
OUT = Path(__file__).resolve().parents[1] / 'docs' / 'Portfolio_Functionality.docx'

"""
Generate a Word (.docx) document from the Markdown summary in
`docs/Portfolio_Functionality.md`.

This script will also convert any SVG diagrams in `docs/images/` to PNG
using `cairosvg` (if available) and embed the images into the generated
`.docx` so diagrams appear inside the Word file.

Dependencies:
- python-docx (pip install python-docx)
- cairosvg (pip install cairosvg)  [optional but recommended to include images]
"""

import sys
from pathlib import Path
import glob

MD = Path(__file__).resolve().parents[1] / 'docs' / 'Portfolio_Functionality.md'
OUT = Path(__file__).resolve().parents[1] / 'docs' / 'Portfolio_Functionality.docx'
IMAGES_DIR = Path(__file__).resolve().parents[1] / 'docs' / 'images'

if not MD.exists():
    print("Markdown source not found:", MD)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches
except Exception as e:
    print("Missing dependency 'python-docx'. Install with: pip install python-docx")
    print("Error:", e)
    sys.exit(2)

try:
    import cairosvg
    HAVE_CAIROSVG = True
except Exception:
    HAVE_CAIROSVG = False

text = MD.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for line in lines:
    line = line.rstrip()
    if not line:
        doc.add_paragraph('')
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:].strip())
    elif line.startswith('1. '):
        p = doc.add_paragraph(style='List Number')
        p.add_run(line[3:].strip())
    else:
        doc.add_paragraph(line)

# Convert SVG diagrams to PNG (if cairosvg available) and embed them
png_files = []
if IMAGES_DIR.exists():
    svg_paths = sorted(IMAGES_DIR.glob('*.svg'))
    for svg in svg_paths:
        png = svg.with_suffix('.png')
        try:
            if HAVE_CAIROSVG:
                # Render at a larger width for crispness; cairosvg will preserve aspect
                cairosvg.svg2png(url=str(svg), write_to=str(png), output_width=1200)
                print(f"Converted: {svg.name} -> {png.name}")
                png_files.append(png)
            else:
                print(f"cairosvg not available; skipping conversion of {svg.name}")
        except Exception as e:
            print(f"Failed to convert {svg.name}: {e}")

    # If PNGs were created, append them to the document under a heading
    if png_files:
        doc.add_page_break()
        doc.add_heading('Design Diagrams', level=2)
        for p in png_files:
            try:
                # Insert image (scale to fit page reasonably)
                doc.add_picture(str(p), width=Inches(6))
                caption = doc.add_paragraph(p.name)
                caption.style = doc.styles['Caption'] if 'Caption' in doc.styles else caption.style
            except Exception as e:
                print(f"Failed to insert image {p.name} into docx: {e}")

try:
    doc.save(OUT)
    print(f"Generated Word document: {OUT}")
except Exception as e:
    print("Failed to save docx:", e)
    sys.exit(3)
